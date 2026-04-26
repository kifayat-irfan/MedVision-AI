from fpdf import FPDF
from docx import Document
from io import BytesIO
import datetime
import re

def sanitize_text(text):
    """
    Removes or replaces special Unicode characters that FPDF (Helvetica) 
    cannot handle. This prevents the FPDFUnicodeEncodingException.
    """
    if not text:
        return ""
    
    # Mapping of fancy characters to simple ASCII characters
    replacements = {
        "\u2013": "-", # en-dash
        "\u2014": "-", # em-dash
        "\u201c": '"', # left double quote
        "\u201d": '"', # right double quote
        "\u2018": "'",  # left single quote
        "\u2019": "'",  # right single quote
        "\u2022": "-", # bullet point
        "\u2026": "...", # ellipsis
        "\u2122": "TM", # trademark
        "\u00ae": "(R)", # registered
        "\u00a9": "(C)", # copyright
    }
    
    # Replace specific fancy characters
    for fancy, simple in replacements.items():
        text = text.replace(fancy, simple)
    
    # Remove any other non-latin-1 characters to be absolutely safe
    return text.encode('latin-1', 'replace').decode('latin-1')

def generate_pdf(text, patient, modality):
    # Step 1: Sanitize the text to avoid Unicode errors
    clean_text = sanitize_text(text)
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "MedVision AI Diagnostic Report", ln=True, align="C")
    
    pdf.set_font("Helvetica", "", 12)
    pdf.ln(10)
    
    # Basic patient info
    pdf.cell(0, 10, f"Modality: {modality}", ln=True)
    pdf.cell(0, 10, f"Patient: {patient}", ln=True)
    pdf.ln(5)
    
    # The actual report
    pdf.multi_cell(0, 10, clean_text)
    
    return bytes(pdf.output())

def generate_docx(text, patient, modality):
    doc = Document()
    doc.add_heading("MedVision AI Diagnostic Report", 0)
    doc.add_paragraph(f"Modality: {modality}\nPatient: {patient}")
    doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

def generate_txt(text):
    return text.encode("utf-8")
